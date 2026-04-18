#!/usr/bin/env python3
"""
Сборка tezisy-markdex-konferenciya-2026.docx по требованиям конференции
«Цифровизация экономики и общества: модели, методы и технологии», апрель 2026.

Формат: А4, книжная; поля: верх 5,9 см, низ 6,4 см, лево и право 4,8 см;
Times New Roman 10,5 pt; абзацный отступ 1 см; интервал одинарный; по ширине.
Объём порядка 8 тыс. знаков с пробелами – контролируется при сборке.

Запуск:
  .venv-tezisy/bin/python docs/konferenciya/build_tezisy_docx.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUT_PATH = Path(__file__).resolve().parent / "tezisy-markdex-konferenciya-2026.docx"


def set_run_font(run, *, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)
    run.bold = bold
    r = run._element
    rpr = r.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:asciiTheme"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:cs"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "Times New Roman")


def apply_margins(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(5.9)
    sec.bottom_margin = Cm(6.4)
    sec.left_margin = Cm(4.8)
    sec.right_margin = Cm(4.8)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    r = p.add_run(text)
    set_run_font(r, bold=True)


def add_meta_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(2)
    pf.space_before = Pt(0)
    r = p.add_run(text)
    set_run_font(r)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.first_line_indent = Cm(1)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    r = p.add_run(text)
    set_run_font(r)


def add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(10)
    pf.space_after = Pt(6)
    r = p.add_run(text.upper())
    set_run_font(r, bold=True)


def add_subheading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.first_line_indent = Cm(1)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(8)
    pf.space_after = Pt(3)
    r = p.add_run(text)
    set_run_font(r, bold=True)


def add_keywords(doc: Document, label: str, words: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(6)
    pf.space_after = Pt(0)
    r1 = p.add_run(label)
    set_run_font(r1, bold=True)
    r2 = p.add_run(words)
    set_run_font(r2)


def add_reference_item(doc: Document, n: int, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.left_indent = Cm(0.35)
    pf.first_line_indent = Cm(-0.35)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(3)
    pf.space_before = Pt(0)
    r = p.add_run(f"{n}. {text}")
    set_run_font(r)


def _estimate_chars(doc: Document) -> int:
    """Грубая оценка знаков с пробелами по абзацам."""
    n = 0
    for p in doc.paragraphs:
        for r in p.runs:
            n += len(r.text)
        n += 1
    return n


def build() -> None:
    doc = Document()
    apply_margins(doc)

    add_title(
        doc,
        "Персональная веб-платформа встроенных карточек внешних источников на базе MDX: "
        "модуль агрегирования метаданных и граница клиент–сервер",
    )

    add_meta_line(doc, "Попов Александр Алексеевич,")
    add_meta_line(
        doc,
        "студент, направление подготовки «Бизнес-информатика», Финансовый университет при Правительстве "
        "Российской Федерации, Москва, Россия, адрес электронной почты указывается в подписном варианте тезисов.",
    )
    add_meta_line(doc, "Научный руководитель:")
    add_meta_line(
        doc,
        "фамилия, имя, отчество, учёная степень, учёное звание, должность, Финансовый университет при "
        "Правительстве Российской Федерации, Москва, Россия, адрес электронной почты указывается в подписном "
        "варианте тезисов.",
    )

    add_body(
        doc,
        "Актуальность определяется повседневной работой целевой аудитории – студентов, преподавателей и "
        "аналитиков, которые собирают подборки статей, регламентов и API-документации: при обычном Markdown "
        "метаданные с целевой страницы переносят вручную, а вставка карт и тяжёлых виджетов в тот же файл "
        "ломает предсказуемый SSR учебной или рабочей страницы. Цель – зафиксировать архитектуру и уже "
        "реализованный минимум персональной веб-платформы на MDX: карточка внешнего источника заполняется из "
        "Open Graph и JSON-LD через серверный HTTP-агент с кэшем, а интерактив редактора отделён от серверной "
        "сборки. Задачи – контракт JSON-API, приоритеты полей DTO, устойчивое поведение при сетевых сбоях. "
        "Итог – связка «редактор – сохранение – карточка» без ручного копирования заголовка и превью и без "
        "смешения серверного и клиентского графа модулей. Практический эффект для вуза – экономия времени на "
        "оформление ссылок в конспектах и хрестоматиях и прозрачный контур для согласования с ИБ; новизна – "
        "согласованное в одном решении разведение нормализации метаданных, версионируемого API и границы SSR.",
    )

    add_keywords(
        doc,
        "Ключевые слова: ",
        "веб-платформа, MDX, метаданные, Open Graph, JSON-LD, SSR, карточка ссылки.",
    )

    add_section_heading(doc, "Введение")

    add_body(
        doc,
        "MDX даёт одному исходному файлу и переносимость Markdown, и выразительность JSX: карточка внешней "
        "страницы описывается компонентом с URL и подставляемыми полями заголовка, описания и превью, без "
        "ручной разметки вокруг каждой ссылки. Для читателя конспекта или методического материала это даёт "
        "единообразие оформления и меньше визуального шума по сравнению с голым списком URL.",
    )
    add_body(
        doc,
        "Чтение HTML стороннего сайта из браузера упирается в same-origin; в проекте реализован серверный "
        "агент с парсингом ответа и отдачей нормализованного JSON в редактор. Модули с доступом к window и "
        "к DOM страницы (карты, сложные виджеты) не входят в серверный бандл и подгружаются на клиенте после "
        "гидратации – так сохраняется стабильный SSR первой выдачи и снимается класс ошибок «сервер вызвал "
        "браузерный API».",
    )

    add_section_heading(doc, "Исследование")

    add_body(
        doc,
        "Реализованный минимум включает окно MDX с предпросмотром, вставку URL статьи или API-документации, "
        "запрос к внутреннему endpoint извлечения метаданных с debounce, автозаполнение полей карточки по "
        "Open Graph и JSON-LD и резерв по тексту заголовка документа и метаописанию, если структурированных "
        "данных нет. Агент нормализует URL, ограничивает редиректы, объём тела и таймаут, проверяет URI "
        "изображения, маппит типы Schema.org на DTO с явным приоритетом при коллизии имён, кэширует ответ по "
        "ключу URL с TTL. В редакторе разделены подсказки для Markdown и для JSX; при обрыве сети файл не "
        "ломается – остаётся URL и компактное отображение по домену.",
    )
    add_body(
        doc,
        "Контракт API задаёт коды 200, 400, 404 и 502, JSON с полями DTO, каноническим URL, признаком источника "
        "данных и версией схемы – это упрощает регрессионные тесты и эволюцию без скрытых поломок клиента. "
        "Ограничены глубина парсинга и логирование; чувствительные query-параметры при необходимости маскируются. "
        "Для целевой аудитории вуза выигрыш прямой: меньше ручного труда на конспект и подборку источников, "
        "одинаковые правила внешнего HTTP и кэша – проще согласовать с подразделением ИБ и воспроизвести решение "
        "в другой группе без копирования «магических» настроек из интерфейса.",
    )
    add_body(
        doc,
        "Проверяемость зафиксирована в коде: к контракту ответа привязаны автоматические проверки на "
        "HTML-фикстурах, поверх них – сквозной сценарий вставки URL в редакторе. При смене версии агента, "
        "парсера или HTTP-клиента преподаватель или сопровождающий получают быстрый сигнал о регрессе без "
        "ручного обхода всех карточек в материалах курса; состав зависимостей задаёт lockfile, что важно для "
        "клонирования репозитория в учебной среде.",
    )
    add_body(
        doc,
        "Ограничение: статический разбор не подменяет headless-браузер – сайты, отдающие смысл только после "
        "клиентского JS, остаются вне гарантии полноты превью; предусмотрены лимиты частоты к домену и учёт "
        "robots.txt. В отличие от связки «Markdown-ссылка плюс отдельный плагин превью», здесь карточка и "
        "агрегация – одна осмысленная цепочка в MDX, что снижает риск рассинхрона между текстом и метаданными.",
    )

    add_section_heading(doc, "Заключение")

    add_body(
        doc,
        "Сформулирован и реализован связный контур: MDX-редактор, серверный агент метаданных с кэшем и явной "
        "границей SSR, дающий целевой аудитории экономию времени и предсказуемое оформление внешних ссылок. "
        "Следующие шаги – расширение покрытия микроразметки, офлайн и политика кэша под нагрузку; накопленный "
        "опыт пригоден как эталон архитектуры для учебных внедрений и оценки совокупной стоимости владения.",
    )

    add_section_heading(doc, "Список использованных источников")
    refs = [
        "ГОСТ Р 7.0.100-2018. Библиографическая запись. Библиографическое описание. Общие требования и правила "
        "составления. – М.: стандартинформ, 2019.",
        "MDX: руководство по синтаксису и компиляции [Электронный ресурс]. – URL: https://mdxjs.com/docs/ "
        "(дата обращения: 28.03.2026).",
        "Open Graph protocol [Электронный ресурс]. – URL: https://ogp.me/ (дата обращения: 28.03.2026).",
        "JSON-LD 1.1 [Электронный ресурс]. – URL: https://www.w3.org/TR/json-ld11/ (дата обращения: 28.03.2026).",
        "Micromark: модульный парсер разметки [Электронный ресурс]. – URL: https://github.com/micromark/micromark "
        "(дата обращения: 28.03.2026).",
        "Schema.org [Электронный ресурс]. – URL: https://schema.org/ (дата обращения: 28.03.2026).",
        "MDXEditor [Электронный ресурс]. – URL: https://mdxeditor.dev/ (дата обращения: 28.03.2026).",
        "HTTP Semantics [Электронный ресурс]. – URL: https://www.rfc-editor.org/rfc/rfc9110 (дата обращения: "
        "28.03.2026).",
        "Fetch Living Standard [Электронный ресурс]. – URL: https://fetch.spec.whatwg.org/ (дата обращения: "
        "28.03.2026).",
        "React: Server Components и границы клиента и сервера [Электронный ресурс]. – URL: "
        "https://react.dev/reference/rsc/server-components (дата обращения: 28.03.2026).",
    ]
    for i, ref in enumerate(refs, start=1):
        add_reference_item(doc, i, ref)

    approx = _estimate_chars(doc)
    doc.save(OUT_PATH)
    print(f"Записано: {OUT_PATH} (оценка знаков с пробелами по тексту абзацев: ~{approx})")


if __name__ == "__main__":
    build()
